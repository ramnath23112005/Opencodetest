from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "outputs"
PLOTS_DIR = OUTPUT_DIR / "plots"
REPORT_PATH = BASE_DIR / "Customer_Churn_Report.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    return table


def add_image(doc, path, width=5.5):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.add_paragraph()
    else:
        add_para(doc, f"[Image not found: {path.name}]", size=9)


def build_report():
    doc = Document()

    title = doc.add_heading("Customer Churn Prediction Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_para(doc, "Automated churn analysis for subscription business | Telco Customer Churn Dataset",
             size=10)
    doc.add_paragraph()

    with open(OUTPUT_DIR / "kpi_report.json") as f:
        kpis = json.load(f)
    with open(OUTPUT_DIR / "model_results.json") as f:
        models = json.load(f)

    add_heading(doc, "Executive Summary", 1)
    add_para(doc,
        f"This report analyzes {kpis['Total Customers']} customers to predict churn and identify key risk factors. "
        f"The overall churn rate is {kpis['Churn Rate (%)']}%, resulting in a monthly revenue loss of "
        f"${kpis['Monthly Revenue Loss ($)']:,.2f}. Retained customers show significantly higher lifetime value "
        f"(${kpis['Avg CLV - Retained ($)']:,.2f}) compared to churned customers "
        f"(${kpis['Avg CLV - Churned ($)']:,.2f})."
    )

    add_heading(doc, "Key Performance Indicators", 2)
    kpi_rows = [
        ["Total Customers", f"{kpis['Total Customers']:,}"],
        ["Churn Rate", f"{kpis['Churn Rate (%)']}%"],
        ["Retention Rate", f"{kpis['Retention Rate (%)']}%"],
        ["Monthly Revenue Loss", f"${kpis['Monthly Revenue Loss ($)']:,.2f}"],
        ["Average CLV", f"${kpis['Average CLV ($)']:,.2f}"],
        ["Avg CLV (Churned)", f"${kpis['Avg CLV - Churned ($)']:,.2f}"],
        ["Avg CLV (Retained)", f"${kpis['Avg CLV - Retained ($)']:,.2f}"],
    ]
    add_table(doc, ["Metric", "Value"], kpi_rows)

    doc.add_paragraph()
    add_heading(doc, "Machine Learning Models", 1)
    add_para(doc,
        "Three classification models were trained using 5-fold stratified cross-validation. "
        "Logistic Regression achieved the highest ROC-AUC score."
    )

    model_headers = ["Model", "ROC-AUC", "Accuracy", "Precision", "Recall", "F1 Score"]
    model_rows = []
    best_model = max(models, key=lambda k: models[k]["ROC-AUC"])
    for name, m in models.items():
        label = f"{name} (best)" if name == best_model else name
        model_rows.append([
            label,
            m["ROC-AUC"],
            m["Accuracy"],
            m["Precision"],
            m["Recall"],
            m["F1"],
        ])
    add_table(doc, model_headers, model_rows)

    doc.add_paragraph()
    add_heading(doc, "ROC Curves Comparison", 2)
    add_image(doc, PLOTS_DIR / "roc_curves_comparison.png")

    add_heading(doc, "Confusion Matrices", 2)
    for name in models:
        add_para(doc, name, bold=True)
        add_image(doc, PLOTS_DIR / f"{name}_confusion_matrix.png", width=4.0)

    add_heading(doc, "Feature Importance", 1)
    add_para(doc,
        "Random Forest feature importance reveals the top predictors of churn. "
        "Tenure, MonthlyCharges, and TotalCharges dominate — customers who are newer, "
        "pay higher monthly fees, and have lower total charges are most at risk."
    )
    add_image(doc, PLOTS_DIR / "feature_importance.png")

    add_heading(doc, "SHAP Analysis", 1)
    add_para(doc,
        "SHAP summary plot (XGBoost) confirms that high MonthlyCharges and low tenure "
        "are the strongest drivers pushing customers toward churn. "
        "Fiber optic internet service and electronic check payment method also increase churn probability."
    )
    add_image(doc, PLOTS_DIR / "shap_summary.png")

    add_heading(doc, "Customer Segmentation", 1)
    add_para(doc,
        "K-Means clustering identified 4 distinct customer segments. "
        "Cluster 3 (short tenure, high charges) has the highest churn rate at ~48%, "
        "while Cluster 2 (long tenure, low charges) has the lowest at ~5%."
    )
    add_image(doc, PLOTS_DIR / "customer_segments.png")
    add_image(doc, PLOTS_DIR / "kmeans_elbow.png", width=4.0)

    add_heading(doc, "Key Insights & Recommendations", 1)

    insights = [
        "Month-to-month contracts have the highest churn rate — promote annual/biannual plans.",
        "Customers with Fiber optic internet churn more than DSL — investigate service quality and pricing.",
        "High support ticket volume strongly correlates with churn — prioritize early intervention.",
        "Customers in the first 12 months of tenure are most vulnerable — implement onboarding engagement.",
        "Electronic check payment method correlates with higher churn — incentivize auto-pay.",
        "Cluster 3 (short tenure, high charges) represents the highest-risk segment — target retention offers.",
    ]
    for ins in insights:
        doc.add_paragraph(ins, style="List Bullet")

    add_heading(doc, "SQL Analysis Queries", 1)
    sql_path = BASE_DIR / "sql" / "churn_analysis.sql"
    if sql_path.exists():
        sql_text = sql_path.read_text()
        add_para(doc, "The following SQL queries analyze churn patterns across contract types, "
                       "payment methods, tenure groups, and services:")
        p = doc.add_paragraph()
        run = p.add_run(sql_text[:2000] + ("\n\n-- (truncated — see sql/churn_analysis.sql for full content)"
                                           if len(sql_text) > 2000 else ""))
        run.font.size = Pt(8)
        run.font.name = "Consolas"

    add_heading(doc, "Dashboard References", 1)
    dashboards = [
        "Power BI Pages: Overview (KPIs), Risk Analysis (segments), Behavior Analysis (tenure/charges vs churn), "
        "Service Analysis (support tickets, internet service).",
        "Tableau Elements: Risk heatmap, customer segmentation clusters, funnel analysis.",
    ]
    for d in dashboards:
        doc.add_paragraph(d, style="List Bullet")

    doc.save(REPORT_PATH)
    print(f"Report saved to {REPORT_PATH}")


if __name__ == "__main__":
    build_report()
